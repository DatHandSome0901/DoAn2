from groq import Groq
import faiss, numpy as np
from sentence_transformers import SentenceTransformer
from docx import Document
import os, time
import re
from dotenv import load_dotenv
# ===== CONFIG =====

load_dotenv()
client = Groq(api_key=os.getenv("GROK_API_KEY"))
model = SentenceTransformer('intfloat/multilingual-e5-base')

docs = [l.strip() for l in open('data/all_text.txt', encoding='utf-8') if l.strip()]
index = faiss.read_index('data/faiss.index')


# ===== SAVE TO WORD =====
def save_to_word(text, topic, outdir):
    # làm sạch tên file
    safe_topic = re.sub(r"[^\w\-]", "_", topic)

    filename = f"de_thi_{safe_topic}_{int(time.time())}.docx"
    path = os.path.join(outdir, filename)

    doc = Document()
    doc.add_heading(f"BỘ CÂU HỎI TỰ LUẬN - CHỦ ĐỀ: {topic}", level=1)

    for line in text.split("\n"):
        doc.add_paragraph(line)

    doc.save(path)
    return filename


# ===== GENERATE EXAM =====
def generate_exam(topic="mạng máy tính", n=10, outdir="output"):
    os.makedirs(outdir, exist_ok=True)

    # embedding query
    qe = model.encode([topic]).astype('float32')
    D, I = index.search(qe, 10)
    ctx = "\n".join([docs[i] for i in I[0]])

    prompt = f"""
Bạn là GIẢNG VIÊN đại học chuyên ngành Công nghệ thông tin.
Trả lời theo chuẩn giáo trình học thuật.

Tạo {n} câu hỏi tự luận về chủ đề: {topic}.
Mỗi câu phải có đáp án chi tiết.

FORMAT BẮT BUỘC:

Câu 1: ...
Đáp án: ...

Câu 2: ...
Đáp án: ...

Chỉ tạo nội dung liên quan đến {topic}.
Không thêm giải thích ngoài format.

Viết theo chuẩn giáo trình Kurose & Ross, Tanenbaum.
Sử dụng thuật ngữ học thuật chính xác.
Phải nêu cơ chế routing, AS, BGP, DNS, ISP hierarchy.
Dữ liệu học liệu:
{ctx}
"""


    res = client.chat.completions.create(
        model='llama-3.1-8b-instant',
        messages=[{'role': 'user', 'content': prompt}],
        max_tokens=2000
    )

    text = res.choices[0].message.content.strip()

    # SAVE WORD
    filename = save_to_word(text, topic, outdir)

    return text, filename
